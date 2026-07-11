from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE = Path(r"C:\Users\Neil\Downloads\MT6_Postman_Submission_Template.docx")
OUTPUT = Path(r"C:\laragon\www\athena-app\MT6_Postman_Submission_Part1_ATHENA.docx")


def replace_paragraph_text(paragraph, text: str) -> None:
    """Replace text while preserving the first run's formatting."""
    template_rpr = None
    if paragraph.runs and paragraph.runs[0]._element.rPr is not None:
        template_rpr = deepcopy(paragraph.runs[0]._element.rPr)

    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)

    run = paragraph.add_run(text)
    if template_rpr is not None:
        run._element.insert(0, template_rpr)


def set_cell_text(cell, text: str, size: float = 8.5) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(size)
            if run._element.rPr is not None:
                fonts = run._element.rPr.rFonts
                if fonts is None:
                    fonts = run._element.rPr.get_or_add_rFonts()
                fonts.set(qn("w:ascii"), "Arial")
                fonts.set(qn("w:hAnsi"), "Arial")


doc = Document(SOURCE)

# Fill the capstone title/date metadata while leaving Name and Section blank.
set_cell_text(
    doc.tables[1].cell(0, 0),
    "Capstone/Thesis Title:\nATHENA (Automated Research Management and Monitoring System)",
    size=9,
)
set_cell_text(doc.tables[1].cell(0, 1), "Date:\nJuly 11, 2026", size=9)

# Part 1: select the application's own Laravel endpoints as the primary case.
replace_paragraph_text(
    doc.paragraphs[6],
    "\u2610 External API(s)      \u2612 Own backend endpoints      \u2610 Stand-in public API (explain why below)",
)
replace_paragraph_text(
    doc.paragraphs[7],
    "Stand-in API explanation: Not applicable.",
)

rows = [
    (
        "{{base_url}}/auth/google",
        "GET",
        "Start Google OAuth sign-in through ATHENA",
        "Sent: none\nReturned: 302 redirect to Google",
        "Secure sign-in using an institutional Google account",
    ),
    (
        "{{base_url}}/auth/google/callback",
        "GET",
        "Handle Google's OAuth callback and sign in the user",
        "Sent: code and state query parameters\nReturned: session and dashboard redirect",
        "Completes Google authentication for ATHENA",
    ),
    (
        "{{base_url}}/research-support/chat",
        "POST",
        "Send a research question to the ATHENA assistant",
        "Sent: messages and optional topic_id\nReturned: reply, model, and usage",
        "Provides AI-assisted research guidance",
    ),
    (
        "{{base_url}}/research-support/literature-search",
        "POST",
        "Search academic literature through ATHENA",
        "Sent: query and filters\nReturned: papers, failed sources, and sources",
        "Helps faculty find related studies",
    ),
    (
        "{{base_url}}/research-support/conference-search",
        "POST",
        "Search for relevant conferences",
        "Sent: search query\nReturned: conference results and sources",
        "Supports publication and dissemination planning",
    ),
]

table = doc.tables[2]
for row_index, values in enumerate(rows, start=1):
    for column_index, value in enumerate(values):
        set_cell_text(table.cell(row_index, column_index), value)

doc.save(OUTPUT)
print(OUTPUT)
